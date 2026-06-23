import os
import re
import hashlib
import tempfile
import json
from datetime import datetime
from copy import copy
import streamlit as st
import pypdf
import openpyxl
from docx import Document
from docx.shared import Inches

# 日期规范化
def normalize_date(date_str: str) -> str:
    if not date_str:
        return ""
    date_str = str(date_str).strip()
    for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日", "%Y-%m", "%Y/%m", "%Y.%m", "%Y年%m月"]:
        try:
            if fmt.endswith("%m"):
                return datetime.strptime(date_str, fmt).strftime("%Y-%m-01")
            else:
                return datetime.strptime(date_str, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    year_match = re.search(r"\b(19|20)\d{2}\b", date_str)
    return f"{year_match.group()}-01-01" if year_match else date_str

def extract_text_from_pdf(pdf_file) -> str:
    reader = pypdf.PdfReader(pdf_file)
    return "\n".join(page.extract_text() for page in reader.pages)

# 图片哈希
def get_image_hash(img_bytes: bytes) -> str:
    return hashlib.sha256(img_bytes).hexdigest()

# 图片OCR和分类（使用全局wrapper）
def analyze_image(img_bytes: bytes, img_filename: str):
    wrapper = st.session_state.get("ai_wrapper")
    if wrapper is None:
        raise RuntimeError("AI wrapper not initialized")
    img_hash = get_image_hash(img_bytes)
    cache = st.session_state.vision_cache
    if img_hash in cache:
        return cache[img_hash]

    text, img_type = wrapper.ocr_and_classify(img_bytes, img_filename)
    cache[img_hash] = (text, img_type)
    return text, img_type

def extract_text_from_images(image_files) -> str:
    if not image_files:
        return ""
    all_texts = []
    for img_file in image_files:
        img_bytes = img_file.read()
        text, _ = analyze_image(img_bytes, img_file.name)
        if text:
            all_texts.append(f"【来自图片：{img_file.name}】\n{text}")
        img_file.seek(0)
    return "\n\n".join(all_texts)

def classify_image_type(img_bytes: bytes, img_filename: str) -> str:
    _, img_type = analyze_image(img_bytes, img_filename)
    return img_type

# Excel 填充辅助函数
def get_primary_cell(worksheet, row, col):
    for merged in worksheet.merged_cells.ranges:
        if (merged.min_row <= row <= merged.max_row and
            merged.min_col <= col <= merged.max_col):
            return worksheet.cell(merged.min_row, merged.min_col)
    return worksheet.cell(row, col)

def safe_write(worksheet, row, col, value):
    get_primary_cell(worksheet, row, col).value = value

def normalize_string(s: str) -> str:
    if not s:
        return ""
    s = str(s).strip()
    s = re.sub(r'[：:，,。、；;]', '', s)
    s = re.sub(r'\s+', '', s)
    return s

def find_cell(worksheet, value, exact=False):
    norm_value = normalize_string(value)
    for row in worksheet.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            if exact:
                if normalize_string(cell.value) == norm_value:
                    return cell.row, cell.column
            else:
                if norm_value in normalize_string(cell.value):
                    return cell.row, cell.column
    return None

def find_row_by_keyword(worksheet, keyword):
    norm_keyword = normalize_string(keyword)
    for row in range(1, worksheet.max_row + 1):
        for col in range(1, worksheet.max_column + 1):
            val = worksheet.cell(row, col).value
            if val and norm_keyword in normalize_string(val):
                return row
    return None

def copy_row_style(source_row, target_row, worksheet, max_col):
    for col in range(1, max_col + 1):
        src = worksheet.cell(source_row, col)
        tgt = worksheet.cell(target_row, col)
        if src.has_style:
            tgt.font = copy(src.font)
            tgt.border = copy(src.border)
            tgt.fill = copy(src.fill)
            tgt.number_format = src.number_format
            tgt.alignment = copy(src.alignment)

def clear_row_content(worksheet, row, max_col):
    for col in range(1, max_col + 1):
        primary = get_primary_cell(worksheet, row, col)
        if primary.row == row and primary.column == col:
            primary.value = None

def fill_basic_info(ws, basic_data):
    for key, value in basic_data.items():
        if not value:
            continue
        pos = find_cell(ws, key, exact=False)
        if pos:
            safe_write(ws, pos[0], pos[1] + 1, value)

def fill_education_block(ws, keyword, edu_data, is_undergrad=True):
    if is_undergrad:
        if "毕业院校" in edu_data and edu_data["毕业院校"]:
            safe_write(ws, 14, 4, edu_data["毕业院校"])
        if "专业" in edu_data and edu_data["专业"]:
            safe_write(ws, 15, 4, edu_data["专业"])
    else:
        if "毕业院校" in edu_data and edu_data["毕业院校"]:
            safe_write(ws, 23, 4, edu_data["毕业院校"])
        if "专业" in edu_data and edu_data["专业"]:
            safe_write(ws, 24, 4, edu_data["专业"])

    pos = find_cell(ws, keyword, exact=True)
    if not pos:
        st.warning(f"未找到关键字: {keyword}")
        return
    row_start = pos[0] + 1
    for offset in range(9):
        current_row = row_start + offset
        field_cell = ws.cell(current_row, 1)
        if field_cell.value is None:
            if offset > 10:
                break
            continue
        field_name = normalize_string(field_cell.value)
        for data_key, data_val in edu_data.items():
            if data_key in ["毕业院校", "专业"]:
                continue
            if data_val and normalize_string(data_key) == field_name:
                safe_write(ws, current_row, 2, data_val)
                break

def fill_work_experience(ws, work_list):
    keyword_row = find_row_by_keyword(ws, "工作经历（由近及远，仅限IT相关经历）")
    if not keyword_row:
        st.warning("未找到「工作经历」关键词，跳过填充")
        return

    header_row = keyword_row + 1
    data_start = header_row + 1
    next_keyword_row = find_row_by_keyword(ws, "项目经历（与上述工作经历匹配，仅IT相关经历）")
    if not next_keyword_row:
        next_keyword_row = ws.max_row + 1

    available_rows = next_keyword_row - data_start
    if available_rows <= 0:
        st.warning("工作经历区域没有预留空行，请在模板中增加空行")
        return

    max_col = ws.max_column
    for row in range(data_start, next_keyword_row):
        clear_row_content(ws, row, max_col)

    sorted_work = sorted(work_list, key=lambda x: x.get("开始日期", "1900-01-01"), reverse=True)
    need = len(sorted_work)
    if need > available_rows:
        st.warning(f"工作经历共有 {need} 条，但模板只预留了 {available_rows} 行，超出部分将被忽略。请手动增加模板预留行数。")
        need = available_rows

    for idx in range(need):
        work = sorted_work[idx]
        target_row = data_start + idx
        for col in range(1, max_col + 1):
            header = ws.cell(header_row, col).value
            if not header:
                continue
            header_norm = normalize_string(header)
            if "开始日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(work.get("开始日期", "")))
            elif "结束日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(work.get("结束日期", "")))
            elif "单位名称" in header_norm:
                safe_write(ws, target_row, col, work.get("单位名称", ""))
            elif "岗位" in header_norm:
                safe_write(ws, target_row, col, work.get("岗位", ""))
            elif "是否邮储银行自主研发工作经验" in header_norm:
                safe_write(ws, target_row, col, work.get("是否邮储银行自主研发工作经验", ""))

def fill_project_experience(ws, project_list):
    keyword_row = find_row_by_keyword(ws, "项目经历（与上述工作经历匹配，仅IT相关经历）")
    if not keyword_row:
        st.warning("未找到「项目经历」关键词，跳过填充")
        return

    header_row = keyword_row + 1
    data_start = header_row + 1
    next_keyword_row = find_row_by_keyword(ws, "技术特长")
    if not next_keyword_row:
        next_keyword_row = ws.max_row + 1

    available_rows = next_keyword_row - data_start
    if available_rows <= 0:
        st.warning("项目经历区域没有预留空行，请在模板中增加空行")
        return

    max_col = ws.max_column
    for row in range(data_start, next_keyword_row):
        clear_row_content(ws, row, max_col)

    sorted_proj = sorted(project_list, key=lambda x: x.get("开始日期", "1900-01-01"), reverse=True)
    need = len(sorted_proj)
    if need > available_rows:
        st.warning(f"项目经历共有 {need} 条，但模板只预留了 {available_rows} 行，超出部分将被忽略。请手动增加模板预留行数。")
        need = available_rows

    for idx in range(need):
        proj = sorted_proj[idx]
        target_row = data_start + idx
        for col in range(1, max_col + 1):
            header = ws.cell(header_row, col).value
            if not header:
                continue
            header_norm = normalize_string(header)
            if "开始日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(proj.get("开始日期", "")))
            elif "结束日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(proj.get("结束日期", "")))
            elif "项目名称" in header_norm:
                safe_write(ws, target_row, col, proj.get("项目名称", ""))
            elif "项目描述" in header_norm:
                safe_write(ws, target_row, col, proj.get("项目描述", ""))
            elif "项目角色" in header_norm:
                safe_write(ws, target_row, col, proj.get("项目角色", ""))
            elif "是否邮储银行自主研发工作经验" in header_norm:
                safe_write(ws, target_row, col, proj.get("是否邮储银行自主研发工作经验", ""))

def fill_template(template_path, output_path, ai_data):
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active
    fill_basic_info(ws, ai_data.get("basic", {}))
    if ai_data.get("education", {}).get("undergraduate"):
        fill_education_block(ws, "本科学历", ai_data["education"]["undergraduate"], is_undergrad=True)
    if ai_data.get("education", {}).get("postgraduate"):
        fill_education_block(ws, "研究生学历", ai_data["education"]["postgraduate"], is_undergrad=False)
    if ai_data.get("work_experience"):
        fill_work_experience(ws, ai_data["work_experience"])
    if ai_data.get("project_experience"):
        fill_project_experience(ws, ai_data["project_experience"])
    wb.save(output_path)

# Word 图片填充
def fill_word_with_images(word_template_bytes, image_classification, new_title=None):
    doc = Document(io.BytesIO(word_template_bytes))
    if new_title:
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            original_font = None
            if first_para.runs:
                original_font = first_para.runs[0].font
            for run in first_para.runs:
                run.clear()
            new_run = first_para.add_run(new_title)
            if original_font:
                new_run.font.name = original_font.name
                new_run.font.size = original_font.size
                new_run.font.bold = original_font.bold
                new_run.font.italic = original_font.italic
                new_run.font.underline = original_font.underline
                if original_font.color and original_font.color.rgb:
                    new_run.font.color.rgb = original_font.color.rgb
    for title, (img_bytes, _) in image_classification.items():
        found = False
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if title in cell.text:
                        if row_idx + 1 < len(table.rows):
                            target_cell = table.cell(row_idx + 1, col_idx)
                        else:
                            target_cell = cell
                        for para in target_cell.paragraphs:
                            drawings = para._element.xpath('.//w:drawing')
                            for draw in drawings:
                                draw.getparent().remove(draw)
                        if target_cell.paragraphs:
                            para = target_cell.paragraphs[0]
                        else:
                            para = target_cell.add_paragraph()
                        run = para.add_run()
                        img_stream = io.BytesIO(img_bytes)
                        run.add_picture(img_stream, width=Inches(5.0))
                        found = True
                        break
                if found:
                    break
            if found:
                break
        if not found:
            st.warning(f"未在表格中找到标题: {title}，请检查模板中的文字是否完全匹配")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
